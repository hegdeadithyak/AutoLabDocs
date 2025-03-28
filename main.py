import json
import io
from pygments import highlight
from pygments.lexers import PythonLexer
from pygments.formatters import ImageFormatter
from docx import Document
from docx.shared import Inches

notebook_path = "FPGA.ipynb"

doc = Document()
doc.add_heading("Notebook Code Cells", level=1)

with open(notebook_path, 'r', encoding='utf-8') as f:
    notebook = json.load(f)

for idx, cell in enumerate(notebook.get('cells', [])):
    if cell.get('cell_type') == 'code':
        code = ''.join(cell.get('source', []))
        if code.strip():
            img_data = highlight(
                code,
                PythonLexer(),
                ImageFormatter(font_name="DejaVu Sans Mono", line_numbers=True)
            )
            image_stream = io.BytesIO(img_data)
            doc.add_picture(image_stream, width=Inches(6))
            doc.add_paragraph("")

doc_stream = io.BytesIO()
doc.save(doc_stream)
doc_stream.seek(0)
output_docx = "Notebook_Code_Cells.docx"
doc.save(output_docx)
print(f"Saved DOCX file: {output_docx}")

