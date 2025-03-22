import os
from docx import Document
from docx.shared import Inches

# Create a new Document
doc = Document()
doc.add_heading("Notebook Code Cells", level=1)

# Loop through the generated PNG images and add them to the document.
# Here, we assume that the images are named as 'code_cell_0.png', 'code_cell_1.png', etc.
idx = 0
while True:
    image_filename = f"code_cell_{idx}.png"
    if not os.path.exists(image_filename):
        break  # Stop if the file doesn't exist
    # Optionally, add a caption or heading for the image
    doc.add_paragraph(f"Code Cell {idx}:")
    # Insert the image (adjust width as needed)
    doc.add_picture(image_filename, width=Inches(6))
    doc.add_paragraph("")  # Add an empty paragraph for spacing
    idx += 1

# Save the document to a file
output_docx = "Notebook_Code_Cells.docx"
doc.save(output_docx)
print(f"Saved DOCX file: {output_docx}")
